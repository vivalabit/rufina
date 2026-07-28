from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = (
    Path(__file__).resolve().parents[1]
    / "apps"
    / "api"
    / "app"
    / "templates"
    / "cover_letter"
    / "standard"
    / "standard-cover-letter.docx"
)
FONT_NAME = "Times New Roman"
BODY_SIZE = 11


def set_run_font(run, *, size: float = BODY_SIZE, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_NAME)


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.15,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_text_paragraph(
    document: Document,
    text: str,
    *,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.15,
    size: float = BODY_SIZE,
    bold: bool = False,
):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(
        paragraph,
        before=before,
        after=after,
        line_spacing=line_spacing,
    )
    set_run_font(paragraph.add_run(text), size=size, bold=bold)
    return paragraph


def build_template() -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.right_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    recipient = document.add_paragraph()
    set_paragraph_spacing(recipient, after=16, line_spacing=1.1)
    for index, line in enumerate(
        (
            "Company Name",
            "Recruiter Name",
            "Street, Number",
            "Postal code, City, Country",
        )
    ):
        if index:
            break_run = recipient.add_run()
            set_run_font(break_run)
            break_run.add_break()
        set_run_font(recipient.add_run(line))

    add_text_paragraph(
        document,
        "City of living, Date",
        after=18,
        line_spacing=1.1,
    )
    subject = add_text_paragraph(
        document,
        "Application for the Position",
        after=18,
        line_spacing=1.1,
        size=16,
        bold=True,
    )
    subject.paragraph_format.keep_with_next = True

    greeting = add_text_paragraph(
        document,
        "Dear Hiring Team,",
        after=12,
        line_spacing=1.15,
    )
    greeting.paragraph_format.keep_with_next = True

    add_text_paragraph(
        document,
        (
            "I am writing to express my interest in this position and explain "
            "why the opportunity is a strong match for my background and motivation."
        ),
        after=10,
    )
    add_text_paragraph(
        document,
        (
            "Throughout my experience, I have built reliable production systems, "
            "scalable services, and automated workflows for real-world applications. "
            "My work has focused on robust architecture, dependable delivery, and "
            "reducing manual effort through thoughtful engineering and automation."
        ),
        after=10,
    )
    add_text_paragraph(
        document,
        (
            "I would welcome the opportunity to bring this experience to the team "
            "and contribute to the role's priorities while continuing to grow."
        ),
        after=16,
    )

    closing = add_text_paragraph(
        document,
        "Kind regards,",
        after=16,
        line_spacing=1.1,
    )
    closing.paragraph_format.keep_with_next = True
    add_text_paragraph(
        document,
        "Candidate Name",
        after=0,
        line_spacing=1.1,
    )

    for paragraph in document.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.widow_control = True

    document.core_properties.title = "Standard Cover Letter"
    document.core_properties.subject = "Bundled cover-letter template"
    document.core_properties.author = "Rufina"
    document.core_properties.keywords = "cover letter, application"

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_template()
