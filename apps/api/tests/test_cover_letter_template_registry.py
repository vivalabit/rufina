from collections.abc import Generator
from io import BytesIO
import struct

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

import app.models.applications  # noqa: F401
from app.core.database import Base, get_db
from app.main import app
from app.services.cover_letter_blocks import extract_cover_letter_blocks_from_docx
from app.services.cover_letter_template_registry import (
    bundled_cover_letter_template_id,
    ensure_bundled_cover_letter_template,
    is_bundled_cover_letter_template_id,
)
from app.services.document_export import (
    build_document_from_template,
    ensure_cover_letter_date_replacement,
)


def test_bundled_cover_letter_template_has_fixed_editable_structure() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine)

    with testing_session_local() as db:
        template = ensure_bundled_cover_letter_template(db)
        blocks = extract_cover_letter_blocks_from_docx(template.content)

    assert template.id == bundled_cover_letter_template_id()
    assert template.name == "Standard cover letter"
    assert template.file_name == "standard-cover-letter.docx"
    assert is_bundled_cover_letter_template_id(template.id) is True
    document = Document(BytesIO(template.content))
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert {
        run.font.name
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text
    } == {"Times New Roman"}
    subject_run = next(
        paragraph.runs[0]
        for paragraph in document.paragraphs
        if paragraph.text == "Application for the Position"
    )
    assert subject_run.font.size.pt == 16
    assert [block["type"] for block in blocks] == [
        "protected",
        "letterDate",
        "subject",
        "greeting",
        "body",
        "body",
        "body",
        "closing",
        "candidateName",
    ]
    assert [
        span["original"]
        for span in blocks[0]["spans"]
        if span["type"] == "text" and span["editable"]
    ] == [
        "Company Name",
        "Recruiter Name",
        "Street, Number",
        "Postal code, City, Country",
    ]


def test_bundled_cover_letter_template_has_real_png_thumbnail() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session_local() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    client = TestClient(app)
    try:
        templates = client.get("/documents/templates/library")
        template_id = templates.json()[0]["id"]
        response = client.get(
            f"/documents/templates/{template_id}/thumbnail",
        )
    finally:
        app.dependency_overrides.clear()
        engine.dispose()

    assert templates.status_code == 200
    assert response.status_code == 200
    assert response.content.startswith(b"\x89PNG\r\n\x1a\n")
    assert response.headers["content-type"] == "image/png"
    assert response.headers["cache-control"] == "private, max-age=300"
    assert response.headers["content-disposition"] == (
        'inline; filename="cover-letter-template-thumbnail.png"'
    )
    assert response.headers["x-rufina-template-id"] == template_id
    width, height = struct.unpack(">II", response.content[16:24])
    assert (width, height) == (360, 640)


def test_bundled_cover_letter_id_is_owner_specific() -> None:
    assert bundled_cover_letter_template_id("owner-a") != bundled_cover_letter_template_id(
        "owner-b"
    )


def test_bundled_template_renders_complete_authoritative_header_and_signature() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session_local = sessionmaker(bind=engine)
    with testing_session_local() as db:
        template = ensure_bundled_cover_letter_template(db)
        template_content = template.content

    enriched = ensure_cover_letter_date_replacement(
        template_content=template_content,
        content='{"replacements":[]}',
        generation_date="2026-07-28",
        language="English",
        vacancy={
            "title": "Practical Trainee - Python Developer (m/f/d)",
            "company": "Sonova",
            "location": "Stäfa, Switzerland",
        },
        profile={
            "name": "Eduard Ishchenko",
            "location": "Männedorf, Zürich, Switzerland",
        },
        recipient_name="Anna Baus",
        official_company_name="Sonova AG",
        recipient_address_line="Laubisrütistrasse 28, 8712 Stäfa, Switzerland",
    )
    rendered = build_document_from_template(
        template_content=template_content,
        content=enriched,
        document_type="cover_letter",
    )
    paragraphs = [paragraph.text for paragraph in Document(BytesIO(rendered)).paragraphs]

    assert paragraphs[:4] == [
        "Sonova AG\nAnna Baus\nLaubisrütistrasse 28\n8712 Stäfa, Switzerland",
        "Männedorf, July 28, 2026",
        "Application for Practical Trainee - Python Developer (m/f/d)",
        "Dear Anna Baus,",
    ]
    assert paragraphs[-2:] == ["Kind regards,", "Eduard Ishchenko"]
