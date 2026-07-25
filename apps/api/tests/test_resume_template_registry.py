from collections.abc import Generator
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.database import Base, get_db
from app.main import app
from app.models.documents import DocumentTemplateRecord
from app.services.document_analysis import analyze_docx_source
from app.services.resume_template_registry import (
    bundled_resume_template_content,
    list_bundled_resume_templates,
    materialize_bundled_resume_template,
)


def test_registry_contains_only_the_three_bundled_resume_templates() -> None:
    templates = list_bundled_resume_templates()

    assert [template.id for template in templates] == [
        "classic_single",
        "modern_single",
        "modern_two_column",
    ]
    assert [template.columns for template in templates] == [1, 1, 2]
    assert len({template.id for template in templates}) == 3


def test_bundled_templates_are_valid_docx_assets() -> None:
    contents = [
        bundled_resume_template_content(template.id)
        for template in list_bundled_resume_templates()
    ]

    assert len(set(contents)) == 3
    for content in contents:
        document = Document(BytesIO(content))
        analysis = analyze_docx_source(content, "tailored_resume")
        assert document.paragraphs
        assert analysis.structure_error == ""
        assert analysis.preflight_report()["supported"] is True
        assert "CANDIDATE NAME" in "\n".join(
            paragraph.text for paragraph in document.paragraphs
        )


def test_resume_template_api_lists_registry_and_rejects_custom_uploads() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session = sessionmaker(
        bind=engine,
        autoflush=False,
        autocommit=False,
    )

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    try:
        client = TestClient(app)
        listed = client.get("/documents/resume-templates")
        uploaded = client.post(
            "/documents/templates",
            json={
                "type": "tailored_resume",
                "name": "User resume",
                "fileName": "user-resume.docx",
                "dataUrl": "data:application/octet-stream;base64,UEs=",
            },
        )

        assert listed.status_code == 200
        assert [item["id"] for item in listed.json()] == [
            "classic_single",
            "modern_single",
            "modern_two_column",
        ]
        assert uploaded.status_code == 422
        assert "Custom resume DOCX templates are not supported" in (
            uploaded.json()["detail"]
        )
        with testing_session() as db:
            assert db.scalar(
                select(func.count()).select_from(DocumentTemplateRecord)
            ) == 0
    finally:
        app.dependency_overrides.clear()
        engine.dispose()


def test_registry_selection_materializes_an_owner_scoped_snapshot() -> None:
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    testing_session = sessionmaker(bind=engine)
    try:
        with testing_session() as db:
            first = materialize_bundled_resume_template(
                db,
                "modern_two_column",
            )
            db.commit()
            second = materialize_bundled_resume_template(
                db,
                "modern_two_column",
            )

            assert first.id == second.id
            assert first.type == "tailored_resume"
            assert first.name == "Modern two-column"
            assert first.content == bundled_resume_template_content(
                "modern_two_column"
            )
            assert db.scalar(
                select(func.count()).select_from(DocumentTemplateRecord)
            ) == 1
    finally:
        engine.dispose()
