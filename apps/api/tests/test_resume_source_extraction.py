from io import BytesIO
import shutil

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from app.services import resume_source_extraction as extraction
from app.services.resume_source_extraction import (
    ResumeSourceExtractionError,
    extract_docx_source_fragments,
    extract_pdf_source_fragments,
)


def document_bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def text_layer_pdf_bytes(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): writer._add_object(font)}
            )
        }
    )
    stream = DecodedStreamObject()
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(
        f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode("latin-1")
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def test_extracts_one_column_docx_into_ordered_paragraph_fragments() -> None:
    document = Document()
    document.add_heading("Ada Lovelace", level=1)
    document.add_paragraph("Platform Engineer")
    document.add_paragraph("Built reliable Python services.")

    result = extract_docx_source_fragments(document_bytes(document))

    assert result.source_format == "docx"
    assert result.layout == "one_column"
    assert result.used_ocr is False
    assert [fragment.text for fragment in result.fragments] == [
        "Ada Lovelace",
        "Platform Engineer",
        "Built reliable Python services.",
    ]
    assert [fragment.order for fragment in result.fragments] == [0, 1, 2]
    assert all(fragment.extraction_method == "docx" for fragment in result.fragments)


def test_extracts_two_column_docx_table_in_visual_column_order() -> None:
    document = Document()
    document.add_paragraph("Ada Lovelace")
    table = document.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.text = "SKILLS"
    left.add_paragraph("Python")
    left.add_paragraph("FastAPI")
    right.text = "EXPERIENCE"
    right.add_paragraph("Platform Engineer")
    right.add_paragraph("Built reliable services.")

    result = extract_docx_source_fragments(document_bytes(document))

    assert result.layout == "two_column"
    assert [fragment.text for fragment in result.fragments] == [
        "Ada Lovelace",
        "SKILLS",
        "Python",
        "FastAPI",
        "EXPERIENCE",
        "Platform Engineer",
        "Built reliable services.",
    ]
    assert [fragment.column_index for fragment in result.fragments] == [
        None,
        0,
        0,
        0,
        1,
        1,
        1,
    ]
    assert all(
        fragment.kind == "table_cell"
        for fragment in result.fragments[1:]
    )


def test_poppler_layout_parser_restores_two_column_pdf_reading_order() -> None:
    layout = b"""<?xml version="1.0" encoding="UTF-8"?>
    <doc xmlns="http://www.w3.org/1999/xhtml">
      <page width="1000" height="1000">
        <flow><block>
          <line xMin="100" yMin="40" xMax="900" yMax="70">
            <word>Ada</word><word>Lovelace</word>
          </line>
          <line xMin="50" yMin="180" xMax="300" yMax="210">
            <word>SKILLS</word>
          </line>
          <line xMin="50" yMin="330" xMax="300" yMax="360">
            <word>Python</word>
          </line>
          <line xMin="50" yMin="480" xMax="300" yMax="510">
            <word>FastAPI</word>
          </line>
          <line xMin="400" yMin="160" xMax="950" yMax="190">
            <word>EXPERIENCE</word>
          </line>
          <line xMin="400" yMin="280" xMax="950" yMax="310">
            <word>Platform</word><word>Engineer</word>
          </line>
          <line xMin="400" yMin="480" xMax="950" yMax="510">
            <word>Built</word><word>services</word>
          </line>
        </block></flow>
      </page>
    </doc>"""

    pages = extraction._parse_pdf_bbox_xml(layout)
    lines, page_layout = extraction._order_pdf_page_lines(list(pages[0].lines))

    assert page_layout == "two_column"
    assert [line.text for line in lines] == [
        "Ada Lovelace",
        "SKILLS",
        "Python",
        "FastAPI",
        "EXPERIENCE",
        "Platform Engineer",
        "Built services",
    ]
    assert [line.column_index for line in lines] == [None, 0, 0, 0, 1, 1, 1]
    assert all(0 <= line.x0 < line.x1 <= 1 for line in lines)


@pytest.mark.skipif(
    shutil.which("pdftotext") is None,
    reason="Poppler pdftotext is not installed",
)
def test_extracts_pdf_text_layer_without_ocr() -> None:
    result = extract_pdf_source_fragments(
        text_layer_pdf_bytes("Ada Lovelace Platform Engineer with Python experience")
    )

    assert result.source_format == "pdf"
    assert result.page_count == 1
    assert result.used_ocr is False
    assert result.layout == "one_column"
    assert result.text == "Ada Lovelace Platform Engineer with Python experience"
    assert result.fragments[0].extraction_method == "pdf_text"


def test_scanned_pdf_page_uses_ocr_fragments(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    empty_page = extraction._PdfPage(
        page_number=1,
        width=612,
        height=792,
        lines=(),
    )
    ocr_lines = extraction._parse_tesseract_tsv(
        "\n".join(
            [
                "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext",
                "1\t1\t0\t0\t0\t0\t0\t0\t1000\t1400\t-1\t",
                "5\t1\t1\t1\t1\t1\t100\t100\t120\t40\t96\tAda",
                "5\t1\t1\t1\t1\t2\t240\t100\t180\t40\t95\tLovelace",
                "5\t1\t1\t1\t2\t1\t100\t180\t210\t40\t94\tEngineer",
            ]
        ),
        page_number=1,
    )
    monkeypatch.setattr(
        extraction,
        "_extract_pdf_text_pages",
        lambda _input, _workdir: [empty_page],
    )
    monkeypatch.setattr(
        extraction,
        "_extract_pdf_ocr_page",
        lambda _input, _workdir, *, page, ocr_languages: ocr_lines,
    )

    result = extract_pdf_source_fragments(b"%PDF-scanned")

    assert result.used_ocr is True
    assert result.page_count == 1
    assert result.layout == "one_column"
    assert [fragment.text for fragment in result.fragments] == [
        "Ada Lovelace",
        "Engineer",
    ]
    assert all(
        fragment.extraction_method == "pdf_ocr"
        for fragment in result.fragments
    )
    assert all(fragment.bbox is not None for fragment in result.fragments)


def test_scanned_pdf_requires_ocr_runtime_when_page_has_no_text(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    page = extraction._PdfPage(
        page_number=1,
        width=612,
        height=792,
        lines=(),
    )
    input_path = tmp_path / "resume.pdf"
    input_path.write_bytes(b"%PDF-scanned")
    monkeypatch.setattr(extraction.shutil, "which", lambda _name: None)

    with pytest.raises(ResumeSourceExtractionError, match="Tesseract OCR"):
        extraction._extract_pdf_ocr_page(
            input_path,
            tmp_path,
            page=page,
            ocr_languages="eng",
        )
