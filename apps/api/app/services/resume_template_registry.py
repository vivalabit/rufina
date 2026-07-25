from __future__ import annotations

import hashlib
from dataclasses import dataclass
from functools import lru_cache
from io import BytesIO
from typing import Literal

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.core.identity import get_bound_owner_id
from app.models.documents import DocumentTemplateRecord, utc_now


ResumeTemplateId = Literal[
    "classic_single",
    "modern_single",
    "modern_two_column",
]
BUNDLED_RESUME_TEMPLATE_VERSION = 1


@dataclass(frozen=True)
class BundledResumeTemplate:
    id: ResumeTemplateId
    name: str
    description: str
    layout: Literal["single_column", "two_column"]
    columns: Literal[1, 2]
    file_name: str


BUNDLED_RESUME_TEMPLATES: tuple[BundledResumeTemplate, ...] = (
    BundledResumeTemplate(
        id="classic_single",
        name="Classic",
        description="Traditional single-column layout optimized for ATS parsing.",
        layout="single_column",
        columns=1,
        file_name="classic-single.docx",
    ),
    BundledResumeTemplate(
        id="modern_single",
        name="Modern",
        description="Contemporary single-column layout with restrained visual hierarchy.",
        layout="single_column",
        columns=1,
        file_name="modern-single.docx",
    ),
    BundledResumeTemplate(
        id="modern_two_column",
        name="Modern two-column",
        description="Two-column layout with a compact skills rail and primary career timeline.",
        layout="two_column",
        columns=2,
        file_name="modern-two-column.docx",
    ),
)
_TEMPLATES_BY_ID = {template.id: template for template in BUNDLED_RESUME_TEMPLATES}


def list_bundled_resume_templates() -> tuple[BundledResumeTemplate, ...]:
    return BUNDLED_RESUME_TEMPLATES


def is_bundled_resume_template_id(template_id: str) -> bool:
    return template_id in _TEMPLATES_BY_ID


def get_bundled_resume_template(template_id: str) -> BundledResumeTemplate:
    try:
        return _TEMPLATES_BY_ID[template_id]
    except KeyError as exc:
        raise KeyError(f"Unknown bundled resume template: {template_id}") from exc


def materialize_bundled_resume_template(
    db: Session,
    template_id: str,
) -> DocumentTemplateRecord:
    template = get_bundled_resume_template(template_id)
    owner_id = get_bound_owner_id()
    record_id = hashlib.sha256(
        (
            f"{owner_id}\0{template.id}\0"
            f"{BUNDLED_RESUME_TEMPLATE_VERSION}"
        ).encode()
    ).hexdigest()[:32]
    existing = db.get(DocumentTemplateRecord, record_id)
    if existing is not None:
        return existing

    content = bundled_resume_template_content(template.id)
    timestamp = utc_now()
    record = DocumentTemplateRecord(
        id=record_id,
        type="tailored_resume",
        name=template.name,
        file_name=template.file_name,
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        content_sha256=hashlib.sha256(content).hexdigest(),
        content=content,
        extracted_text=template.description,
        created_at=timestamp,
        updated_at=timestamp,
    )
    db.add(record)
    db.flush()
    db.refresh(record)
    return record


@lru_cache(maxsize=3)
def bundled_resume_template_content(template_id: ResumeTemplateId) -> bytes:
    template = get_bundled_resume_template(template_id)
    document = Document()
    configure_page(document)
    if template.id == "classic_single":
        build_classic_single(document)
    elif template.id == "modern_single":
        build_modern_single(document)
    else:
        build_modern_two_column(document)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)


def add_identity(document: Document, *, accent: RGBColor | None = None) -> None:
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("CANDIDATE NAME")
    run.bold = True
    run.font.size = Pt(20)
    if accent is not None:
        run.font.color.rgb = accent
    contacts = document.add_paragraph(
        "City · email@example.com · +00 000 000 000 · linkedin.com/in/profile"
    )
    contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(document: Document, heading: str, placeholder: str) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(heading.upper())
    run.bold = True
    run.font.size = Pt(11)
    document.add_paragraph(placeholder)


def build_classic_single(document: Document) -> None:
    add_identity(document)
    add_section(document, "Summary", "Evidence-backed professional summary.")
    add_section(
        document,
        "Experience",
        "Role · Company · Dates\n• Evidence-backed accomplishment and impact.",
    )
    add_section(document, "Skills", "Relevant skill · Relevant skill")
    add_section(document, "Education", "Credential · Institution · Dates")


def build_modern_single(document: Document) -> None:
    accent = RGBColor(0x1F, 0x5A, 0x7A)
    add_identity(document, accent=accent)
    for heading, placeholder in (
        ("Profile", "Focused, evidence-backed value proposition."),
        (
            "Selected experience",
            "Role · Company · Dates\n• Measured achievement delivered by a clear method.",
        ),
        ("Core capabilities", "Capability · Capability · Capability"),
        ("Education", "Credential · Institution · Dates"),
    ):
        add_section(document, heading, placeholder)


def build_modern_two_column(document: Document) -> None:
    accent = RGBColor(0x24, 0x3B, 0x53)
    add_identity(document, accent=accent)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(2.1)
    right.width = Inches(4.8)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_cell_section(left, "Skills", "Relevant capability\nRelevant capability")
    add_cell_section(left, "Education", "Credential\nInstitution · Dates")
    add_cell_section(right, "Profile", "Focused, evidence-backed value proposition.")
    add_cell_section(
        right,
        "Experience",
        "Role · Company · Dates\n• Measured achievement delivered by a clear method.",
    )


def add_cell_section(cell, heading: str, placeholder: str) -> None:
    title = cell.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    run = title.add_run(heading.upper())
    run.bold = True
    run.font.size = Pt(10)
    cell.add_paragraph(placeholder)


__all__ = [
    "BUNDLED_RESUME_TEMPLATE_VERSION",
    "BUNDLED_RESUME_TEMPLATES",
    "BundledResumeTemplate",
    "ResumeTemplateId",
    "bundled_resume_template_content",
    "get_bundled_resume_template",
    "is_bundled_resume_template_id",
    "list_bundled_resume_templates",
    "materialize_bundled_resume_template",
]
