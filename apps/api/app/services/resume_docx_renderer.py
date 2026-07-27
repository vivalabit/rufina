from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pydantic import ValidationError

from app.models.resume import FinalResume
from app.models.resume_templates import ResumeTemplateDesignTokens
from app.services.resume_pdf_renderer import resume_section_titles, skill_group_view


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class ResumeDocxRenderError(RuntimeError):
    pass


def render_final_resume_docx(
    final_resume_json: dict[str, object],
    *,
    design: ResumeTemplateDesignTokens,
) -> bytes:
    """Render a validated FinalResume as an editable Word document."""
    try:
        resume = FinalResume.model_validate(final_resume_json)
        document = Document()
        configure_document(document, design)
        add_identity(document, resume, design)
        add_resume_sections(document, resume, design)
        document.core_properties.title = f"{resume.basics.full_name} - Resume"
        document.core_properties.subject = "Tailored resume"
        document.core_properties.author = resume.basics.full_name
        output = BytesIO()
        document.save(output)
        return output.getvalue()
    except ValidationError as exc:
        raise ResumeDocxRenderError(
            "FinalResume JSON failed schema validation"
        ) from exc
    except ResumeDocxRenderError:
        raise
    except Exception as exc:
        raise ResumeDocxRenderError("Resume DOCX rendering failed") from exc


def configure_document(
    document: Document,
    design: ResumeTemplateDesignTokens,
) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(design.page_margins.top)
    section.right_margin = Mm(design.page_margins.right)
    section.bottom_margin = Mm(design.page_margins.bottom)
    section.left_margin = Mm(design.page_margins.left)

    font_size = 10 * design.font_scale
    normal = document.styles["Normal"]
    normal.font.name = design.font_family
    normal.font.size = Pt(font_size)
    normal.paragraph_format.space_after = Pt(density_spacing(design.density))
    normal.paragraph_format.line_spacing = 1.04
    set_style_font(normal, design.font_family)

    for style_name in ("Title", "Subtitle"):
        set_style_font(document.styles[style_name], design.font_family)

    section_style = document.styles.add_style(
        "Resume Section",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    section_style.base_style = normal
    section_style.font.name = design.font_family
    section_style.font.size = Pt(10.5 * design.font_scale)
    section_style.font.bold = True
    section_style.font.color.rgb = hex_to_rgb(design.accent_color)
    section_style.font.all_caps = True
    section_style.paragraph_format.space_before = Pt(
        density_section_spacing(design.density)
    )
    section_style.paragraph_format.space_after = Pt(3)
    section_style.paragraph_format.keep_with_next = True
    set_style_font(section_style, design.font_family)

    item_style = document.styles.add_style(
        "Resume Item",
        WD_STYLE_TYPE.PARAGRAPH,
    )
    item_style.base_style = normal
    item_style.paragraph_format.space_before = Pt(2)
    item_style.paragraph_format.space_after = Pt(0.5)
    item_style.paragraph_format.keep_with_next = True

    bullet_style = document.styles["List Bullet"]
    bullet_style.base_style = normal
    bullet_style.paragraph_format.left_indent = Mm(4.5)
    bullet_style.paragraph_format.first_line_indent = Mm(-3)
    bullet_style.paragraph_format.space_after = Pt(
        1 if design.density == "compact" else 2
    )
    set_style_font(bullet_style, design.font_family)


def add_identity(
    document: Document,
    resume: FinalResume,
    design: ResumeTemplateDesignTokens,
) -> None:
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    name.paragraph_format.keep_with_next = True
    run = name.add_run(resume.basics.full_name)
    run.bold = True
    run.font.name = design.font_family
    run.font.size = Pt(24 * design.font_scale)
    run.font.color.rgb = hex_to_rgb(design.accent_color)

    if resume.basics.headline:
        headline = document.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.paragraph_format.space_after = Pt(3)
        headline.paragraph_format.keep_with_next = True
        run = headline.add_run(resume.basics.headline)
        run.bold = True
        run.font.size = Pt(11 * design.font_scale)

    contact_values = [
        ("email", resume.basics.email),
        ("phone", resume.basics.phone),
        ("text", resume.basics.location),
        ("text", resume.basics.work_authorization),
        ("link", resume.basics.linkedin),
        ("link", resume.basics.github),
        ("link", resume.basics.portfolio),
    ]
    visible_contacts = [(kind, value) for kind, value in contact_values if value]
    if not visible_contacts:
        return
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    for index, (kind, value) in enumerate(visible_contacts):
        if index:
            contact.add_run("  |  ")
        if kind == "email":
            add_hyperlink(contact, value, f"mailto:{value}")
        elif kind == "phone":
            add_hyperlink(contact, value, f"tel:{value}")
        elif kind == "link":
            add_hyperlink(contact, value, web_target(value))
        else:
            contact.add_run(value)


def add_resume_sections(
    document: Document,
    resume: FinalResume,
    design: ResumeTemplateDesignTokens,
) -> None:
    titles = resume_section_titles(resume)
    for section in resume.section_order:
        if section == "summary" and resume.summary:
            add_section_heading(document, titles[section], design)
            document.add_paragraph(resume.summary.text)
        elif section == "experience" and resume.experiences:
            add_section_heading(document, titles[section], design)
            for experience in resume.experiences:
                add_item_heading(
                    document,
                    experience.title,
                    experience.period,
                    design,
                )
                add_item_meta(
                    document,
                    joined([experience.company, experience.location]),
                )
                add_bullets(
                    document,
                    [bullet.text for bullet in experience.bullets],
                )
        elif section == "skills" and resume.skills:
            add_section_heading(document, titles[section], design)
            for group in skill_group_view(resume):
                category = str(group["category"])
                names = [str(name) for name in group["names"]]
                paragraph = document.add_paragraph()
                if category:
                    run = paragraph.add_run(f"{category}: ")
                    run.bold = True
                paragraph.add_run(", ".join(names))
        elif section == "education" and resume.education:
            add_section_heading(document, titles[section], design)
            for education in resume.education:
                qualification = joined(
                    [education.credential, education.field_of_study],
                    separator=", ",
                )
                period = date_range(education.start_date, education.end_date)
                add_item_heading(document, qualification, period, design)
                add_item_meta(
                    document,
                    joined([education.institution, education.location]),
                )
                add_bullets(
                    document,
                    [detail.text for detail in education.details],
                )
        elif section == "projects" and resume.projects:
            add_section_heading(document, titles[section], design)
            for project in resume.projects:
                add_item_heading(document, project.name, project.role, design)
                if project.url:
                    paragraph = document.add_paragraph(style="Resume Item")
                    add_hyperlink(
                        paragraph,
                        project.url,
                        web_target(project.url),
                    )
                add_bullets(
                    document,
                    [bullet.text for bullet in project.bullets],
                )
        elif section == "certifications" and resume.certifications:
            add_section_heading(document, titles[section], design)
            for certification in resume.certifications:
                dates = certification.issued_on
                if certification.expires_on:
                    dates = joined(
                        [
                            dates,
                            f"{titles['expires']} {certification.expires_on}",
                        ],
                    )
                add_item_heading(
                    document,
                    certification.name,
                    dates,
                    design,
                )
                add_item_meta(document, certification.issuer)
        elif section == "languages" and resume.languages:
            add_section_heading(document, titles[section], design)
            for language in resume.languages:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(language.name)
                run.bold = True
                paragraph.add_run(f": {language.proficiency}")
        elif section == "additional" and resume.additional_sections:
            add_section_heading(document, titles[section], design)
            for additional in resume.additional_sections:
                paragraph = document.add_paragraph(style="Resume Item")
                run = paragraph.add_run(additional.title)
                run.bold = True
                add_bullets(
                    document,
                    [item.text for item in additional.items],
                )


def add_section_heading(
    document: Document,
    title: str,
    design: ResumeTemplateDesignTokens,
) -> None:
    paragraph = document.add_paragraph(title, style="Resume Section")
    if design.heading_style in {"underlined", "accent-rule"}:
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), design.accent_color.removeprefix("#"))
        borders.append(bottom)
        properties.append(borders)


def add_item_heading(
    document: Document,
    primary: str,
    trailing: str,
    design: ResumeTemplateDesignTokens,
) -> None:
    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(
            210
            - design.page_margins.left
            - design.page_margins.right
            - 1
        ),
        WD_TAB_ALIGNMENT.RIGHT,
    )
    run = paragraph.add_run(primary)
    run.bold = True
    if trailing:
        paragraph.add_run(f"\t{trailing}")


def add_item_meta(document: Document, value: str) -> None:
    if not value:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(value)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_bullets(document: Document, values: list[str]) -> None:
    for value in values:
        document.add_paragraph(value, style="List Bullet")


def add_hyperlink(paragraph, text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_style_font(style, font_family: str) -> None:
    style.font.name = font_family
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)


def hex_to_rgb(value: str) -> RGBColor:
    normalized = value.removeprefix("#")
    return RGBColor(
        int(normalized[0:2], 16),
        int(normalized[2:4], 16),
        int(normalized[4:6], 16),
    )


def web_target(value: str) -> str:
    return value if "://" in value else f"https://{value}"


def joined(values: list[str], *, separator: str = " | ") -> str:
    return separator.join(value for value in values if value)


def date_range(start: str, end: str) -> str:
    return joined([start, end], separator=" - ")


def density_spacing(density: str) -> float:
    return {"compact": 1.5, "standard": 2.5, "comfortable": 4}[density]


def density_section_spacing(density: str) -> float:
    return {"compact": 5, "standard": 7, "comfortable": 9}[density]


__all__ = [
    "DOCX_CONTENT_TYPE",
    "ResumeDocxRenderError",
    "render_final_resume_docx",
]
